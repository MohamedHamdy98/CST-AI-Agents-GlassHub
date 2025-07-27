from docx import Document

with open(r'H:\GlassHub\cst_agent\brainstorm\examples\controls.json', 'r', encoding='utf-8') as f:
    text = f.read()

document = Document()
document.add_heading('Raw JSON Content', level=1)
document.add_paragraph(text)
document.save('raw_json_output.docx')
