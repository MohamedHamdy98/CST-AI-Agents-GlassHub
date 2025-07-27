import os, logging, glob, sys, httpx, json
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from utils.schemas import FileURLs, FileURLsRag
from docx import Document
from dotenv import load_dotenv
from rag.knowledge_ingestion import ingest_company_knowledge
from utils.helper_functions import (chunk_pages, estimate_chunk_size,save_temp_file, 
                                    convert_docx_to_pdf, extract_pages_from_pdf, init_oss_bucket,
                                    upload_to_alibaba_oss_static, download_files_from_cloud_storage,
                                    retrieve_full_knowledge_from_docx, load_documents,
                                    process_all_formatted_results)


# Load environment variables from .env file
load_dotenv()

# Logger setup
logger = logging.getLogger("report_router")
handler = logging.StreamHandler(sys.stdout)
handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
handler.stream.reconfigure(encoding='utf-8')  
logger.addHandler(handler)
logger.setLevel(logging.INFO)

router = APIRouter(prefix="/api/v1", tags=["admin"])


# Alibaba cloud Connection
ACCESS_KEY_ID = os.getenv("OSS_ACCESS_KEY_ID")
ACCESS_KEY_SECRET = os.getenv("OSS_ACCESS_KEY_SECRET")
ENDPOINT = os.getenv("OSS_ENDPOINT")
BUCKET_NAME = os.getenv("OSS_BUCKET")


# Qwen3 LLM Endpoint
QWEN3_ENDPOINT = os.getenv('QWEN3_ENDPOINT')
QWEN3_ENDPOINT_CHAT = os.getenv('QWEN3_ENDPOINT_CHAT')



@router.post("/extract_terms", description="Extract terms from Word file using LLM")
def extract_terms(
    word_file: UploadFile = File(...),
    name_word_file: str = Form(...),
    language: str = Form(..., description="Arabic, English"),
    max_tokens: int = Form(512),
    thinking: bool = Form(False),
    timeout: int = Form(180)
):
    word_path = ""
    response_path = "./database/llm_raw_outputs"
    try:
        logger.info("📥 Endpoint /extract_clauses_sync called.")
        logger.info(f"Received file: {word_file.filename}, name_word_file: {name_word_file}, max_tokens: {max_tokens}, thinking: {thinking}, timeout: {timeout}")

        if not word_file.filename.endswith(".docx"):
            logger.warning("Uploaded file is not a .docx")
            return JSONResponse(status_code=400, content={"error": "Uploaded file must be a .docx"})

        os.makedirs("tmp", exist_ok=True)
        word_path = save_temp_file(word_file, f"tmp/{word_file.filename}")
        pdf_path = word_path.replace(".docx", ".pdf")
        logger.info(f"Saved DOCX to {word_path} and will convert to PDF at {pdf_path}")

        # ✅ Convert DOCX to PDF
        convert_docx_to_pdf(word_path, pdf_path)
        logger.info("✅ DOCX converted to PDF.")

        # ✅ Extract pages from PDF
        pages = extract_pages_from_pdf(pdf_path)
        chunk_size = estimate_chunk_size(pages, max_tokens=max_tokens)
        chunks = chunk_pages(pages, chunk_size=chunk_size)
        logger.info(f"Extracted {len(pages)} pages into {len(chunks)} chunks.")

        # ✅ Prepare prompts
        prompts = [
            f"""Your task is to extract (التعليمات والبنود والمطلوب من المستخدم ان يقوم به او ما يجب ان يتجنبه ) from raw text:\n\n{chunk}. \n the language is {language}"""
            for chunk in chunks
        ]

        all_responses = ""
        for idx, prompt in enumerate(prompts):
            try:
                logger.info(f"⏳ Sending prompt for chunk {idx+1}/{len(chunks)}")
                headers = {
                    "accept": "application/json",
                    "Content-Type": "application/x-www-form-urlencoded"
                }
                data = {
                    "prompt": prompt,
                    "name_word_file": name_word_file,
                    "max_tokens": str(max_tokens),
                    "thinking": str(thinking).lower()
                }

                response = httpx.post(QWEN3_ENDPOINT, headers=headers, data=data, timeout=timeout)

                if response.status_code != 200:
                    logger.error(f"❌ LLM API failed on chunk {idx+1}: {response.status_code} - {response.text}")
                    raise Exception(f"LLM API failed: {response.status_code} - {response.text}")

                llm_response = response.json().get("response", "")
                logger.info(f"✅ Received response for chunk {idx+1}")
                all_responses += llm_response + "," +"\n"

            except Exception as e:
                logger.exception(f"❌ Error in chunk {idx+1}: {str(e)}")

        # ✅ Save to docx
        llm_outputs = "./database/llm_raw_outputs"
        os.makedirs(llm_outputs, exist_ok=True)
        final_docx_path = f"./database/llm_raw_outputs/{name_word_file}_response.docx"
        doc = Document()
        doc.add_paragraph(all_responses.strip())
        doc.save(final_docx_path)
        logger.info(f"✅ Saved LLM output to {final_docx_path}")

        # ✅ Upload to Alibaba cloud storage
        blob_name = f"{name_word_file}_response.docx"
        bucket = init_oss_bucket(ACCESS_KEY_ID, ACCESS_KEY_SECRET, ENDPOINT, BUCKET_NAME)
        url = upload_to_alibaba_oss_static(bucket, final_docx_path, f"cst_rag/{blob_name}")
        logger.info(f"📤 Uploaded file to OOS Alibab. URL: {url}")

        return {"llm_response": all_responses.strip(), "url": url}

    except Exception as e:
        logger.exception(f"🔥 Unexpected error: {str(e)}")
        return JSONResponse(status_code=500, content={"error": str(e)})

    finally:
        for file_path in glob.glob(os.path.join(response_path, "*")):
            try:
                os.remove(file_path)
                print(f"Deleted: {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
        for file in [word_path, pdf_path]:
            try:
                if os.path.exists(file):
                    os.remove(file)
                    logger.info(f"🧹 Deleted temp file: {file}")
            except Exception as cleanup_err:
                logger.warning(f"⚠️ Could not delete temp file {file}: {cleanup_err}")


@router.post("/generate_controls")
def generate_controls(
    file_urls_json: FileURLs
):
    """ 
        {
        "files": [
            {
                "url": "https://example.com/file1.docx",
                "name_file": "File1_Controls"
            },
            {
                "url": "https://example.com/file2.docx",
                "name_file": "File2_Controls"
            }
            ]
        }

    """
    DOCX_DIRECTORY = "./database/controls"
    results = []

    try:
        if not file_urls_json or not file_urls_json.files:
            logger.exception("Input list is required")
            raise HTTPException(status_code=400, detail="files are required")

        for entry in file_urls_json.files:
            url = entry.url
            name_file = entry.name_file
            OUTPUT_DOCX = f"./database/controls/{name_file}.docx"

            logger.info(f"Processing: {name_file} from {url}")

            # 1. Download file
            download_files_from_cloud_storage([url], DOCX_DIRECTORY)

            # 2. Load documents
            documents = load_documents()

            # 3. Retrieve knowledge
            retrieve_doc = retrieve_full_knowledge_from_docx(documents)

            # 4. Process results
            js_result = process_all_formatted_results(retrieve_doc)

            # 5. Save to docx
            document = Document()
            document.add_heading("Controls", level=1)
            document.add_paragraph(json.dumps(js_result, ensure_ascii=False, indent=2))
            document.save(OUTPUT_DOCX)
            logger.info(f"✅ Saved DOCX: {OUTPUT_DOCX}")

            # 6. Upload to OSS
            bucket = init_oss_bucket(ACCESS_KEY_ID, ACCESS_KEY_SECRET, ENDPOINT, BUCKET_NAME)
            oss_path = f"controls_results/{name_file}.docx"
            url_uploaded = upload_to_alibaba_oss_static(bucket, OUTPUT_DOCX, oss_path)

            # 7. Append result
            results.append({
                "url": url_uploaded,
                "name_file": name_file,
                "result": js_result,
                "uploaded_doc": oss_path
            })

            # 8. Clean individual output
            if os.path.exists(OUTPUT_DOCX):
                os.remove(OUTPUT_DOCX)

        return {
            "success": True,
            "files_processed": results
        }

    except Exception as e:
        logger.exception("❌ Failed to process or upload.")
        return {
            "message": f"Failed to process or upload: {str(e)}",
            "success": False
        }

    finally:
        # 🧹 Cleanup any leftover docx files
        try:
            for file in os.listdir(DOCX_DIRECTORY):
                if file.endswith(".docx"):
                    os.remove(os.path.join(DOCX_DIRECTORY, file))
                    logger.info(f"🧹 Cleaned up input: {file}")
        except Exception as cleanup_err:
            logger.warning(f"⚠️ Cleanup failed: {cleanup_err}")


# For RAG System
@router.post("/create_rag_system")
def create_rag_system(file_urls_json: FileURLsRag):

    DOCX_DIRECTORY = "./database/glasshub_files"

    try:
        # ✅ Validate input
        if not file_urls_json:
            logger.exception("file_urls is required")
            raise HTTPException(status_code=400, detail="file_urls is required")

        file_urls = file_urls_json.urls
        logger.info(f"The URLs are {file_urls}")

        # 🧠 Ingest knowledge from files
        ingest_company_knowledge(file_urls)
        logger.info("📥 Files are downloaded and processed successfully.")

        return {
            "message": "RAG system created and files processed.",
            "success": True
        }

    except Exception as e:
        logger.exception("❌ Failed to download or process files.")
        return {
            "message": f"Failed to download or process files: {str(e)}.",
            "success": False
        }

    finally:
        # 🧹 Clean up temporary .docx files from local storage
        try:
            for file in os.listdir(DOCX_DIRECTORY):
                if file.endswith(".docx"):
                    full_path = os.path.join(DOCX_DIRECTORY, file)
                    if os.path.exists(full_path):
                        os.remove(full_path)
                        logger.info(f"🧹 Deleted temp file: {full_path}")
        except Exception as cleanup_err:
            logger.warning(f"⚠️ Cleanup failed: {cleanup_err}")

