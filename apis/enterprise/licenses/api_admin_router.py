import os, logging, glob, sys, httpx, json
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from typing import List, Union
from utils.schemas import FileURLs, FileURLsRag
from docx import Document
from dotenv import load_dotenv
from rag.knowledge_ingestion import ingest_company_knowledge
from utils.logs import setup_logger
from utils.helper_functions import (chunk_pages, estimate_chunk_size,save_temp_file, 
                                    convert_docx_to_pdf, extract_pages_from_pdf, init_oss_bucket,
                                    upload_to_alibaba_oss_static, download_files_from_cloud_storage,
                                    retrieve_full_knowledge_from_docx, load_documents,
                                    process_all_formatted_results, process_all_formatted_results_no_llm)


# Load environment variables from .env file
load_dotenv()

# Logger setup


logger = logger = setup_logger(__name__)
print(f"🔥 DEBUG: Module {__name__} logger setup complete!")
print(f"🔥 DEBUG: Module {__name__} logger setup complete!")

router = APIRouter(prefix="/api/v1/enterprise/licenses", tags=["Enterprise Licenses Admin"])


# Alibaba cloud Connection
ACCESS_KEY_ID = os.getenv("OSS_ACCESS_KEY_ID")
ACCESS_KEY_SECRET = os.getenv("OSS_ACCESS_KEY_SECRET")
ENDPOINT = os.getenv("OSS_ENDPOINT")
BUCKET_NAME = os.getenv("OSS_BUCKET")


# Qwen3 LLM Endpoint
QWEN3_ENDPOINT = os.getenv('QWEN3_ENDPOINT')
QWEN3_ENDPOINT_CHAT = os.getenv('QWEN3_ENDPOINT_CHAT')


@router.post("/extract_terms", description="Extract terms from multiple Word files using LLM")
def extract_terms_enterprise_licenses(
    word_file: List[UploadFile] = File(..., description="Upload one or more Word files"),
    name_word_file: Union[List[str], str] = Form(..., description="Matching names for the uploaded files"),
    language: str = Form(..., description="Arabic or English"),
    max_tokens: int = Form(512),
    thinking: bool = Form(False),
    timeout: int = Form(180)
):
    response_path = "./database/llm_raw_outputs/enterprise/licenses"
    os.makedirs(response_path, exist_ok=True)
    all_results = []

    # Normalize names into a list
    if isinstance(name_word_file, list):
        if len(name_word_file) == 1 and "," in name_word_file[0]:
            # User sent a single comma-separated string
            names_list = [name.strip() for name in name_word_file[0].split(",")]
        else:
            # User sent multiple form fields
            names_list = [name.strip() for name in name_word_file]
    else:
        # User sent a plain string (rare case)
        names_list = [name.strip() for name in name_word_file.split(",")]

    try:
        logger.info("📥 Endpoint /extract_terms called.")
        logger.info(f"Received {len(word_file)} files with names: {names_list}")

        if len(word_file) != len(names_list):
            return JSONResponse(status_code=400, content={"error": "Number of files and names must match"})

        for file, file_name in zip(word_file, names_list):
            logger.info(f"Processing file: {file.filename} as {file_name}")

            # ✅ Validate extension
            if not file.filename.endswith(".docx"):
                logger.warning(f"❌ Skipping {file.filename} (not .docx)")
                continue

            # ✅ Save temp file
            os.makedirs("tmp", exist_ok=True)
            word_path = save_temp_file(file, f"tmp/{file.filename}")
            pdf_path = word_path.replace(".docx", ".pdf")
            logger.info(f"Saved DOCX to {word_path}, converting to {pdf_path}")

            # ✅ Convert DOCX → PDF → Extract Pages → Chunk
            convert_docx_to_pdf(word_path, pdf_path)
            pages = extract_pages_from_pdf(pdf_path)
            chunk_size = estimate_chunk_size(pages, max_tokens=max_tokens)
            chunks = chunk_pages(pages, chunk_size=chunk_size)
            logger.info(f"{file.filename}: {len(pages)} pages → {len(chunks)} chunks")

            # ✅ Build prompts and collect responses
            all_responses = ""
            for idx, chunk in enumerate(chunks):
                prompt = (
                    f"Your task is to extract (التعليمات والبنود والمطلوب من المستخدم ان يقوم به او ما يجب ان يتجنبه ) "
                    f"from raw text:\n\n{chunk}\nThe language is {language}"
                )
                try:
                    logger.info(f"⏳ Sending prompt {idx+1}/{len(chunks)} for {file.filename}")
                    headers = {"accept": "application/json", "Content-Type": "application/x-www-form-urlencoded"}
                    data = {
                        "prompt": prompt,
                        "name_word_file": file_name,
                        "max_tokens": str(max_tokens),
                        "thinking": str(thinking).lower(),
                    }
                    response = httpx.post(QWEN3_ENDPOINT, headers=headers, data=data, timeout=timeout)

                    if response.status_code != 200:
                        raise Exception(f"LLM API failed: {response.status_code} - {response.text}")

                    llm_response = response.json().get("response", "")
                    all_responses += llm_response + "\n"
                except Exception as e:
                    logger.exception(f"❌ Error processing chunk {idx+1} for {file.filename}: {str(e)}")

            # ✅ Save combined output
            final_docx_path = os.path.join(response_path, f"{file_name}.docx")
            doc = Document()
            doc.add_paragraph(all_responses.strip())
            doc.save(final_docx_path)
            logger.info(f"✅ Saved output to {final_docx_path}")

            # ✅ Upload to Alibaba OSS
            blob_name = f"{file_name}_response.docx"
            bucket = init_oss_bucket(ACCESS_KEY_ID, ACCESS_KEY_SECRET, ENDPOINT, BUCKET_NAME)
            oss_path = f"cst/cst_extract_terms/enterprise/licenses{blob_name}"
            url = upload_to_alibaba_oss_static(bucket, final_docx_path, oss_path)
            logger.info(f"📤 Uploaded {blob_name} to Alibaba Cloud → {url}")

            all_results.append({"file_name": file_name, "llm_response": all_responses.strip(), "url": url, "uploaded_doc": oss_path})

            # ✅ Cleanup temp files
            for f in [word_path, pdf_path]:
                if os.path.exists(f):
                    os.remove(f)
                    logger.info(f"🧹 Deleted temp file: {f}")

        return {"results": all_results}

    except Exception as e:
        logger.exception(f"🔥 Unexpected error: {str(e)}")
        return JSONResponse(status_code=500, content={"error": str(e)})

    finally:
        # Clear old responses
        for file_path in glob.glob(os.path.join(response_path, "*")):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.warning(f"⚠️ Could not delete {file_path}: {e}")



@router.post("/generate_controls")
def generate_controls_enterprise_licenses(
    file_urls_json: FileURLs
):
    """ 
        {
        "files": [
            {
                "url": "https://example.com/file1.docx",
                "name_file": "File1_Controls"
            },
            {
                "url": "https://example.com/file2.docx",
                "name_file": "File2_Controls"
            }
            ]
        }

    """
    DOCX_DIRECTORY = "./database/controls/enterprise/licenses"
    results = []

    try:
        if not file_urls_json or not file_urls_json.files:
            logger.exception("Input list is required")
            raise HTTPException(status_code=400, detail="files are required")

        for entry in file_urls_json.files:
            url = entry.url
            name_file = entry.name_file
            OUTPUT_DOCX = f"./database/controls/enterprise/{name_file}.docx"

            logger.info(f"Processing: {name_file} from {url}")

            # 1. Download file
            download_files_from_cloud_storage([url], DOCX_DIRECTORY)

            # 2. Load documents
            documents = load_documents(docx_path=DOCX_DIRECTORY)

            # 3. Retrieve knowledge
            retrieve_doc = retrieve_full_knowledge_from_docx(documents)

            # 4. Process results
            js_result = process_all_formatted_results(retrieve_doc)

            # 5. Save to docx
            document = Document()
            document.add_heading("Controls", level=1)
            document.add_paragraph(json.dumps(js_result, ensure_ascii=False, indent=2))
            document.save(OUTPUT_DOCX)
            logger.info(f"✅ Saved DOCX: {OUTPUT_DOCX}")

            # 6. Upload to OSS
            bucket = init_oss_bucket(ACCESS_KEY_ID, ACCESS_KEY_SECRET, ENDPOINT, BUCKET_NAME)
            oss_path = f"cst/cst_controls/controls_results/enterprise/license/{name_file}.docx"
            url_uploaded = upload_to_alibaba_oss_static(bucket, OUTPUT_DOCX, oss_path)

            # 7. Append result
            results.append({
                "url": url_uploaded,
                "name_file": name_file,
                "result": js_result,
                "uploaded_doc": oss_path
            })

            # 8. Clean individual output
            if os.path.exists(OUTPUT_DOCX):
                os.remove(OUTPUT_DOCX)

        return {
            "success": True,
            "files_processed": results
        }

    except Exception as e:
        logger.exception("❌ Failed to process or upload.")
        return {
            "message": f"Failed to process or upload: {str(e)}",
            "success": False
        }

    finally:
        # 🧹 Cleanup any leftover docx files
        try:
            for file in os.listdir(DOCX_DIRECTORY):
                if file.endswith(".docx"):
                    os.remove(os.path.join(DOCX_DIRECTORY, file))
                    logger.info(f"Cleaned up input: {file}")
        except Exception as cleanup_err:
            logger.warning(f"⚠️ Cleanup failed: {cleanup_err}")



# For RAG System
@router.post("/create_rag_system")
def create_rag_system(file_urls_json: FileURLsRag):

    DOCX_DIRECTORY = "./database/glasshub_files/enterprise/licenses"
    path_upload = "cst/cst_rag/enterprise/licenses"
    try:
        # ✅ Validate input
        if not file_urls_json:
            logger.exception("file_urls is required")
            raise HTTPException(status_code=400, detail="file_urls is required")

        file_urls = file_urls_json.urls
        logger.info(f"The URLs are {file_urls}")

        # 🧠 Ingest knowledge from files
        ingest_company_knowledge(file_urls, DOCX_DIRECTORY, path_upload)
        logger.info("📥 Files are downloaded and processed successfully.")

        return {
            "message": "RAG system created and files processed.",
            "success": True,
            "path_faiss":f"https://glasshub-files-staging.oss-me-central-1.aliyuncs.com/{path_upload}/index.faiss",
            "path_pkl":f"https://glasshub-files-staging.oss-me-central-1.aliyuncs.com/{path_upload}/index.pkl"
        }

    except Exception as e:
        logger.exception("❌ Failed to download or process files.")
        return {
            "message": f"Failed to download or process files: {str(e)}.",
            "success": False
        }

    finally:
        # 🧹 Clean up temporary .docx files from local storage
        try:
            for file in os.listdir(DOCX_DIRECTORY):
                if file.endswith(".docx"):
                    full_path = os.path.join(DOCX_DIRECTORY, file)
                    if os.path.exists(full_path):
                        os.remove(full_path)
                        logger.info(f"🧹 Deleted temp file: {full_path}")
        except Exception as cleanup_err:
            logger.warning(f"⚠️ Cleanup failed: {cleanup_err}")

